from __future__ import annotations

import shutil
from datetime import datetime, timedelta, timezone
from pathlib import Path
from typing import Optional

from fastapi import APIRouter, Depends, File, Form, HTTPException, UploadFile
from fastapi.responses import PlainTextResponse
from sqlalchemy import func, select
from sqlalchemy.orm import Session, selectinload

from app import __version__
from app.auth import get_optional_user_id, verify_api_key
from app.cache import clone_source_from_cache, extract_youtube_video_id, find_cached_source
from app.config import get_settings
from app.db import get_db
from app.jobs import run_in_background
from app.limits import enforce_daily_source_limit
from app.llm.qa import answer_question
from app.llm.summarize import summarize_segments
from app.llm_settings_store import llm_status, save_llm_overrides
from app.share import make_share_slug
from app.models import Ask, Segment, Source, Summary
from app.pipeline import (
    process_article_source,
    process_podcast_source,
    process_text_source,
    process_upload_source,
    process_youtube_source,
    reprocess_source_from_media,
)
from app.ingest.podcast import fetch_rss_episodes
from app.ingest.youtube import expand_youtube_urls
from app.schemas import (
    AskOut,
    AskRequest,
    AskResponse,
    HealthResponse,
    JobStatusOut,
    LlmSettingsUpdate,
    PlaylistCreateRequest,
    PodcastRssCreateRequest,
    PodcastEpisodeCreateRequest,
    ArticleCreateRequest,
    QuotaOut,
    ShareOut,
    ReprocessRequest,
    SourceDetailOut,
    SourceOut,
    SummarizeRequest,
    SummaryOut,
    TextCreateRequest,
    YouTubeCreateRequest,
)

router = APIRouter()
protected = APIRouter(dependencies=[Depends(verify_api_key)])


def _to_source_out(source: Source, segment_count: int | None = None) -> SourceOut:
    if segment_count is None:
        if source.segments is not None:
            segment_count = len(source.segments)
        else:
            segment_count = 0
    tags = []
    try:
        tags = [tag.name for tag in (source.tags or [])]
    except Exception:
        tags = []
    return SourceOut(
        id=source.id,
        user_id=source.user_id,
        source_type=source.source_type,
        title=source.title,
        url=source.url,
        video_id=source.video_id,
        language=source.language,
        status=source.status,
        progress=float(getattr(source, "progress", 0) or 0),
        progress_message=getattr(source, "progress_message", "") or "",
        error=source.error,
        error_code=source.error_code,
        error_hint=source.error_hint,
        duration_seconds=source.duration_seconds,
        transcript_method=source.transcript_method,
        description=getattr(source, "description", None),
        author=getattr(source, "author", None),
        show_title=getattr(source, "show_title", None),
        published_at=getattr(source, "published_at", None),
        created_at=source.created_at,
        updated_at=source.updated_at,
        segment_count=segment_count,
        share_slug=source.share_slug,
        is_public=bool(source.is_public),
        tags=tags,
    )


def _get_owned_source(
    db: Session,
    source_id: int,
    user_id: str,
    *,
    load_segments: bool = False,
    load_summaries: bool = False,
) -> Source:
    stmt = select(Source).where(Source.id == source_id, Source.user_id == user_id)
    options = [selectinload(Source.tags)]
    if load_segments:
        options.append(selectinload(Source.segments))
    if load_summaries:
        options.append(selectinload(Source.summaries))
    if options:
        stmt = stmt.options(*options)
    source = db.scalar(stmt)
    if source is None:
        raise HTTPException(status_code=404, detail="Source not found")
    return source


def _segment_counts(db: Session, source_ids: list[int]) -> dict[int, int]:
    if not source_ids:
        return {}
    rows = db.execute(
        select(Segment.source_id, func.count())
        .where(Segment.source_id.in_(source_ids))
        .group_by(Segment.source_id)
    ).all()
    return {int(source_id): int(count) for source_id, count in rows}


@router.get("/health", response_model=HealthResponse)
def health() -> HealthResponse:
    return HealthResponse(status="ok", version=__version__)


@protected.get("/llm/status")
def get_llm_status() -> dict:
    return llm_status()


@protected.put("/llm/settings")
def update_llm_settings(payload: LlmSettingsUpdate) -> dict:
    save_llm_overrides(payload.model_dump(exclude_none=True))
    return llm_status()


@protected.get("/sources", response_model=list[SourceOut])
def list_sources(
    db: Session = Depends(get_db),
    user_id: str = Depends(get_optional_user_id),
) -> list[SourceOut]:
    sources = list(
        db.scalars(
            select(Source)
            .where(Source.user_id == user_id)
            .options(selectinload(Source.tags))
            .order_by(Source.id.desc())
        ).all()
    )
    counts = _segment_counts(db, [s.id for s in sources])
    return [_to_source_out(s, segment_count=counts.get(s.id, 0)) for s in sources]


@protected.get("/sources/{source_id}", response_model=SourceDetailOut)
def get_source(
    source_id: int,
    db: Session = Depends(get_db),
    user_id: str = Depends(get_optional_user_id),
) -> SourceDetailOut:
    source = _get_owned_source(
        db, source_id, user_id, load_segments=True, load_summaries=True
    )
    base = _to_source_out(source)
    return SourceDetailOut(
        **base.model_dump(),
        segments=source.segments,
        summaries=source.summaries,
    )


@protected.get("/sources/{source_id}/status", response_model=JobStatusOut)
def get_status(
    source_id: int,
    db: Session = Depends(get_db),
    user_id: str = Depends(get_optional_user_id),
) -> JobStatusOut:
    source = _get_owned_source(db, source_id, user_id)
    progress_map = {
        "pending": "Queued",
        "downloading": "Downloading media",
        "transcribing": "Transcribing / extracting text",
        "summarizing": "Generating summary",
        "ready": "Ready",
        "failed": "Failed",
    }
    return JobStatusOut(
        source_id=source.id,
        status=source.status,
        error=source.error,
        error_code=source.error_code,
        error_hint=source.error_hint,
        progress=getattr(source, "progress_message", None) or progress_map.get(source.status, source.status),
        progress_pct=float(getattr(source, "progress", 0) or 0),
        progress_message=getattr(source, "progress_message", "") or "",
    )


@protected.get("/sources/{source_id}/asks", response_model=list[AskOut])
def list_asks(
    source_id: int,
    db: Session = Depends(get_db),
    user_id: str = Depends(get_optional_user_id),
) -> list[AskOut]:
    _get_owned_source(db, source_id, user_id)
    asks = db.scalars(
        select(Ask).where(Ask.source_id == source_id).order_by(Ask.created_at.desc())
    ).all()
    return list(asks)


@protected.post("/sources/youtube", response_model=SourceOut)
def create_youtube_source(
    payload: YouTubeCreateRequest,
    db: Session = Depends(get_db),
    user_id: str = Depends(get_optional_user_id),
) -> SourceOut:
    settings = get_settings()
    enforce_daily_source_limit(db, user_id, settings)

    video_id = extract_youtube_video_id(payload.url)
    cached = find_cached_source(db, video_id, user_id) if video_id else None

    source = Source(
        user_id=user_id,
        source_type="youtube",
        title="YouTube video",
        url=payload.url,
        video_id=video_id,
        language=payload.language,
        status="pending",
    )
    db.add(source)
    db.commit()
    db.refresh(source)

    if cached is not None:
        clone_source_from_cache(db, source, cached)
        db.commit()
        source = _get_owned_source(db, source.id, user_id, load_segments=True)
        return _to_source_out(source)

    run_in_background(process_youtube_source, source.id, auto_summarize=payload.auto_summarize)
    db.refresh(source)
    return _to_source_out(source, segment_count=0)


@protected.post("/sources/text", response_model=SourceOut)
def create_text_source(
    payload: TextCreateRequest,
    db: Session = Depends(get_db),
    user_id: str = Depends(get_optional_user_id),
) -> SourceOut:
    enforce_daily_source_limit(db, user_id, get_settings(), source_type="text")
    source = Source(
        user_id=user_id,
        source_type="text",
        title=payload.title,
        language=payload.language,
        status="pending",
        transcript_method="text",
    )
    db.add(source)
    db.commit()
    db.refresh(source)

    settings = get_settings()
    dest_dir = settings.media_dir / f"source_{source.id}"
    dest_dir.mkdir(parents=True, exist_ok=True)
    dest = dest_dir / "content.txt"
    dest.write_text(payload.text, encoding="utf-8")
    source.file_path = str(dest)
    db.commit()

    run_in_background(
        process_text_source,
        source.id,
        text=payload.text,
        auto_summarize=payload.auto_summarize,
    )
    db.refresh(source)
    return _to_source_out(source, segment_count=0)


@protected.post("/sources/upload", response_model=SourceOut)
async def upload_source(
    file: UploadFile = File(...),
    language: str = Form("pl"),
    title: Optional[str] = Form(None),
    auto_summarize: bool = Form(True),
    db: Session = Depends(get_db),
    user_id: str = Depends(get_optional_user_id),
) -> SourceOut:
    settings = get_settings()
    original_name = file.filename or "upload.bin"
    suffix = Path(original_name).suffix.lower()
    if suffix not in {".pdf", ".txt", ".md", ".docx", ".epub", ".mp3", ".wav", ".m4a", ".m4b", ".webm", ".ogg", ".opus", ".flac", ".aac", ".mp4", ".mkv"}:
        raise HTTPException(status_code=400, detail=f"Unsupported file type: {suffix}")

    if suffix == ".pdf":
        source_type = "pdf"
    elif suffix == ".epub":
        source_type = "book"
    elif suffix == ".docx":
        source_type = "document"
    elif suffix in {".mp3", ".wav", ".m4a", ".m4b", ".webm", ".ogg", ".opus", ".flac", ".aac", ".mp4", ".mkv"}:
        source_type = "audiobook" if suffix == ".m4b" else "audio"
    else:
        source_type = "text"
    enforce_daily_source_limit(db, user_id, settings, source_type=source_type)
    source = Source(
        user_id=user_id,
        source_type=source_type,
        title=title or original_name,
        language=language,
        status="pending",
    )
    db.add(source)
    db.commit()
    db.refresh(source)

    dest_dir = settings.media_dir / f"source_{source.id}"
    dest_dir.mkdir(parents=True, exist_ok=True)
    dest = dest_dir / original_name
    with dest.open("wb") as out:
        shutil.copyfileobj(file.file, out)
    source.file_path = str(dest)
    db.commit()

    run_in_background(process_upload_source, source.id, auto_summarize=auto_summarize)
    db.refresh(source)
    return _to_source_out(source, segment_count=0)


@protected.post("/sources/{source_id}/reprocess", response_model=SourceOut)
def reprocess_source(
    source_id: int,
    payload: ReprocessRequest,
    db: Session = Depends(get_db),
    user_id: str = Depends(get_optional_user_id),
) -> SourceOut:
    source = _get_owned_source(db, source_id, user_id)
    source.status = "pending"
    source.error = None
    source.error_code = None
    source.error_hint = None
    db.commit()
    run_in_background(
        reprocess_source_from_media,
        source.id,
        prefer_captions=payload.prefer_captions,
        force_asr=payload.force_asr,
        auto_summarize=payload.auto_summarize,
    )
    db.refresh(source)
    return _to_source_out(source, segment_count=_segment_counts(db, [source.id]).get(source.id, 0))


@protected.post("/sources/{source_id}/summarize", response_model=SummaryOut)
def summarize_source(
    source_id: int,
    payload: SummarizeRequest,
    db: Session = Depends(get_db),
    user_id: str = Depends(get_optional_user_id),
) -> SummaryOut:
    source = _get_owned_source(db, source_id, user_id, load_segments=True)
    if source.status != "ready":
        raise HTTPException(status_code=409, detail=f"Source not ready (status={source.status})")

    segs = [(s.start, s.end, s.text) for s in source.segments]
    content = summarize_segments(segs, title=source.title, kind=payload.kind)
    summary = Summary(source_id=source.id, kind=payload.kind, content=content)
    db.add(summary)
    db.commit()
    db.refresh(summary)
    return summary


@protected.post("/ask", response_model=AskResponse)
def ask(
    payload: AskRequest,
    db: Session = Depends(get_db),
    user_id: str = Depends(get_optional_user_id),
) -> AskResponse:
    if payload.source_id is None:
        raise HTTPException(status_code=400, detail="source_id is required")
    source = _get_owned_source(db, payload.source_id, user_id, load_segments=True)
    if source.status != "ready":
        raise HTTPException(status_code=409, detail=f"Source not ready (status={source.status})")

    segs = [(s.start, s.end, s.text) for s in source.segments]
    answer, citations = answer_question(payload.question, segs, title=source.title)
    db.add(Ask(source_id=source.id, question=payload.question, answer=answer))
    db.commit()
    return AskResponse(answer=answer, citations=citations, source_id=source.id)


@protected.delete("/sources/{source_id}")
def delete_source(
    source_id: int,
    db: Session = Depends(get_db),
    user_id: str = Depends(get_optional_user_id),
) -> dict[str, str]:
    source = _get_owned_source(db, source_id, user_id)
    settings = get_settings()
    media_dir = settings.media_dir / f"source_{source.id}"
    db.delete(source)
    db.commit()
    if media_dir.exists():
        shutil.rmtree(media_dir, ignore_errors=True)
    return {"status": "deleted"}




@protected.get("/quota", response_model=QuotaOut)
def get_quota(
    db: Session = Depends(get_db),
    user_id: str = Depends(get_optional_user_id),
) -> QuotaOut:
    settings = get_settings()
    limit = settings.daily_source_limit
    since = datetime.now(timezone.utc) - timedelta(days=1)
    used = db.scalar(
        select(func.count()).select_from(Source).where(Source.user_id == user_id, Source.created_at >= since)
    ) or 0
    remaining = max(0, limit - used) if limit > 0 else 10**9
    return QuotaOut(used=int(used), limit=int(limit), remaining=int(remaining))


@protected.post("/sources/playlist", response_model=list[SourceOut])
def create_playlist(
    payload: PlaylistCreateRequest,
    db: Session = Depends(get_db),
    user_id: str = Depends(get_optional_user_id),
) -> list[SourceOut]:
    settings = get_settings()
    urls = [u.strip() for u in payload.urls if u and u.strip()]
    if not urls:
        raise HTTPException(status_code=400, detail="No URLs provided")
    if len(urls) > 20:
        raise HTTPException(status_code=400, detail="Playlist limited to 20 URLs")

    created: list[SourceOut] = []
    for url in urls:
        enforce_daily_source_limit(db, user_id, settings)
        video_id = extract_youtube_video_id(url)
        cached = find_cached_source(db, video_id, user_id) if video_id else None
        source = Source(
            user_id=user_id,
            source_type="youtube",
            title="YouTube video",
            url=url,
            video_id=video_id,
            language=payload.language,
            status="pending",
        )
        db.add(source)
        db.commit()
        db.refresh(source)
        if cached is not None:
            clone_source_from_cache(db, source, cached)
            db.commit()
            source = _get_owned_source(db, source.id, user_id, load_segments=True)
            created.append(_to_source_out(source))
            continue
        run_in_background(process_youtube_source, source.id, auto_summarize=payload.auto_summarize)
        created.append(_to_source_out(source, segment_count=0))
    return created


@protected.post("/sources/{source_id}/share", response_model=ShareOut)
def share_source(
    source_id: int,
    db: Session = Depends(get_db),
    user_id: str = Depends(get_optional_user_id),
) -> ShareOut:
    source = _get_owned_source(db, source_id, user_id, load_summaries=True)
    if source.status != "ready":
        raise HTTPException(status_code=409, detail="Source not ready")
    if not source.summaries:
        raise HTTPException(status_code=409, detail="No summary to share")
    if not source.share_slug:
        source.share_slug = make_share_slug()
    source.is_public = True
    db.commit()
    settings = get_settings()
    base = (settings.public_base_url or "").rstrip("/")
    share_url = f"{base}/s/{source.share_slug}" if base else f"/s/{source.share_slug}"
    return ShareOut(
        source_id=source.id,
        share_slug=source.share_slug,
        share_url=share_url,
        is_public=True,
    )


@protected.get("/sources/{source_id}/export.md")
def export_markdown(
    source_id: int,
    db: Session = Depends(get_db),
    user_id: str = Depends(get_optional_user_id),
):
    source = _get_owned_source(db, source_id, user_id, load_segments=True, load_summaries=True)
    latest = source.summaries[-1].content if source.summaries else ""
    lines = [
        f"# {source.title or 'Podsumowanie'}",
        "",
        f"- Status: {source.status}",
        f"- Method: {source.transcript_method or '-'}",
        f"- URL: {source.url or '-'}",
        "",
        latest or "_Brak podsumowania_",
        "",
        "## Transkrypt",
        "",
    ]
    for seg in source.segments:
        mm = int(seg.start // 60)
        ss = int(seg.start % 60)
        lines.append(f"- [{mm:02d}:{ss:02d}] {seg.text}")
    return PlainTextResponse("\n".join(lines), media_type="text/markdown; charset=utf-8")






@protected.get("/sources/{source_id}/export.docx")
def export_source_docx(
    source_id: int,
    db: Session = Depends(get_db),
    user_id: str = Depends(get_optional_user_id),
):
    from io import BytesIO
    from docx import Document
    from fastapi.responses import StreamingResponse

    source = _get_owned_source(db, source_id, user_id, load_segments=True, load_summaries=True)
    doc = Document()
    doc.add_heading(source.title or "Source", level=1)
    if source.show_title:
        doc.add_paragraph(f"Show: {source.show_title}")
    if source.author:
        doc.add_paragraph(f"Author: {source.author}")
    for summary in source.summaries:
        doc.add_heading(summary.kind, level=2)
        doc.add_paragraph(summary.content)
    doc.add_heading("Transcript", level=2)
    for seg in source.segments:
        doc.add_paragraph(f"[{int(seg.start)//60:02d}:{int(seg.start)%60:02d}] {seg.text}")
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    filename = f"source-{source.id}.docx"
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@protected.post("/sources/article", response_model=SourceOut)
def create_article_source(
    payload: ArticleCreateRequest,
    db: Session = Depends(get_db),
    user_id: str = Depends(get_optional_user_id),
) -> SourceOut:
    if extract_youtube_video_id(payload.url):
        raise HTTPException(status_code=400, detail="Use /sources/youtube for YouTube URLs")
    enforce_daily_source_limit(db, user_id, get_settings(), source_type="article")
    source = Source(
        user_id=user_id,
        source_type="article",
        title=payload.title or "Article",
        url=payload.url,
        language=payload.language,
        status="pending",
    )
    db.add(source)
    db.commit()
    db.refresh(source)
    run_in_background(process_article_source, source.id, auto_summarize=payload.auto_summarize)
    return _to_source_out(source, segment_count=0)


@protected.post("/sources/podcast", response_model=SourceOut)
def create_podcast_episode(
    payload: PodcastEpisodeCreateRequest,
    db: Session = Depends(get_db),
    user_id: str = Depends(get_optional_user_id),
) -> SourceOut:
    if extract_youtube_video_id(payload.url):
        raise HTTPException(status_code=400, detail="Use /sources/youtube for YouTube URLs")
    enforce_daily_source_limit(db, user_id, get_settings(), source_type="podcast")
    source = Source(
        user_id=user_id,
        source_type="podcast",
        title="Podcast episode",
        url=payload.url,
        language=payload.language,
        status="pending",
    )
    db.add(source)
    db.commit()
    db.refresh(source)
    run_in_background(process_podcast_source, source.id, auto_summarize=payload.auto_summarize)
    return _to_source_out(source, segment_count=0)


@protected.post("/sources/podcast/rss", response_model=list[SourceOut])
def create_podcast_rss(
    payload: PodcastRssCreateRequest,
    db: Session = Depends(get_db),
    user_id: str = Depends(get_optional_user_id),
) -> list[SourceOut]:
    settings = get_settings()
    episodes = fetch_rss_episodes(payload.feed_url, max_episodes=payload.max_episodes)
    if not episodes:
        raise HTTPException(status_code=400, detail="No audio episodes found in RSS feed")
    created: list[SourceOut] = []
    for episode in episodes:
        enforce_daily_source_limit(db, user_id, settings, source_type="podcast")
        source = Source(
            user_id=user_id,
            source_type="podcast",
            title=episode.title,
            url=episode.audio_url,
            language=payload.language,
            status="pending",
            show_title=episode.show_title,
            description=episode.description,
            author=episode.author,
            published_at=episode.published_at,
        )
        db.add(source)
        db.commit()
        db.refresh(source)
        run_in_background(process_podcast_source, source.id, auto_summarize=payload.auto_summarize)
        created.append(_to_source_out(source, segment_count=0))
    return created


@protected.post("/sources/url", response_model=SourceOut)
def create_from_url(
    payload: ArticleCreateRequest,
    db: Session = Depends(get_db),
    user_id: str = Depends(get_optional_user_id),
) -> SourceOut:
    """Smart router: YouTube -> youtube, audio file URL/podcast -> podcast, else article."""
    url = payload.url.strip()
    if extract_youtube_video_id(url):
        yt = YouTubeCreateRequest(url=url, language=payload.language, auto_summarize=payload.auto_summarize)
        return create_youtube_source(yt, db=db, user_id=user_id)
    lower = url.lower()
    if any(lower.endswith(ext) for ext in (".mp3", ".m4a", ".m4b", ".ogg", ".opus", ".wav", ".flac", ".aac")) or "rss" in lower or "/feed" in lower:
        pod = PodcastEpisodeCreateRequest(url=url, language=payload.language, auto_summarize=payload.auto_summarize)
        return create_podcast_episode(pod, db=db, user_id=user_id)
    return create_article_source(payload, db=db, user_id=user_id)


router.include_router(protected)
