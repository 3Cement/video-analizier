from __future__ import annotations

from pathlib import Path

from docx import Document


def extract_docx_text(path: Path) -> str:
    document = Document(str(path))
    parts: list[str] = []
    for paragraph in document.paragraphs:
        text = (paragraph.text or "").strip()
        if text:
            parts.append(text)
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text and cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    joined = "\n\n".join(parts).strip()
    if len(joined) < 40:
        raise RuntimeError("DOCX extraction produced too little content")
    return joined
