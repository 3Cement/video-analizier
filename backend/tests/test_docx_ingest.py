from pathlib import Path

from docx import Document

from app.ingest.docx_ingest import extract_docx_text


def test_extract_docx_text(tmp_path: Path):
    path = tmp_path / "note.docx"
    doc = Document()
    doc.add_paragraph("This document has enough characters for the DOCX extractor threshold in tests.")
    doc.add_paragraph("Second paragraph with more words about audiobooks and articles.")
    doc.save(path)
    text = extract_docx_text(path)
    assert "enough characters" in text
